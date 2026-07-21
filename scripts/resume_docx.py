#!/usr/bin/env python3
import argparse, json, shutil, sys
from pathlib import Path
from docx import Document

def paragraphs(doc):
    """Yield body and table paragraphs once, including merged table cells."""
    seen = set()
    for paragraph in doc.paragraphs:
        seen.add(id(paragraph._p))
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    key = id(paragraph._p)
                    if key not in seen:
                        seen.add(key)
                        yield paragraph

def is_bullet(paragraph):
    text = paragraph.text.strip()
    properties = paragraph._p.pPr
    has_numbering = properties is not None and properties.numPr is not None
    return bool(text) and (
        has_numbering
        or 'list' in paragraph.style.name.lower()
        or text.lstrip().startswith(('•', '-'))
    )

def extract(path):
    doc = Document(path)
    content = list(paragraphs(doc))
    print(json.dumps({
        'text': '\n'.join(p.text.strip() for p in content if p.text.strip()),
        'bullets': [p.text.strip() for p in content if is_bullet(p)],
    }))

def generate(source,target,bullets):
    if not bullets or any(not isinstance(value, str) or not value.strip() for value in bullets):
        raise RuntimeError('Generated bullets must all contain text')
    shutil.copy2(source,target)
    doc = Document(target)
    candidates = [p for p in paragraphs(doc) if is_bullet(p)]
    if not candidates: raise RuntimeError('No bullet paragraphs found in master resume')
    for paragraph,value in zip(candidates,bullets):
        if paragraph.runs:
            paragraph.runs[0].text=value
            for run in paragraph.runs[1:]: run.text=''
        else: paragraph.text=value
    doc.save(target); print(json.dumps({'path':str(Path(target).resolve()),'replaced':min(len(candidates),len(bullets))}))

p=argparse.ArgumentParser(); p.add_argument('command',choices=['extract','generate']); p.add_argument('--source',required=True); p.add_argument('--target'); args=p.parse_args()
try:
    if args.command=='extract': extract(args.source)
    else: generate(args.source,args.target,json.load(sys.stdin)['bullets'])
except Exception as e:
    print(str(e),file=sys.stderr); sys.exit(2)
